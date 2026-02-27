import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import threading
import math
import os

class CSVtoWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("终极 CSV 转 Word 表格排版神器")
        self.root.geometry("500x300")
        self.root.resizable(False, False)

        # UI 元素布局
        self.setup_ui()
        self.csv_path = ""

    def setup_ui(self):
        # 标题
        title_label = tk.Label(self.root, text="CSV 自动排版至 Word 工具", font=("Arial", 16, "bold"))
        title_label.pack(pady=20)

        # 文件选择区域
        file_frame = tk.Frame(self.root)
        file_frame.pack(pady=10)
        
        self.path_label = tk.Label(file_frame, text="尚未选择文件...", fg="gray", width=40, anchor="w")
        self.path_label.pack(side=tk.LEFT, padx=10)
        
        btn_browse = tk.Button(file_frame, text="浏览 CSV", command=self.browse_file)
        btn_browse.pack(side=tk.LEFT)

        # 进度显示
        self.status_var = tk.StringVar()
        self.status_var.set("准备就绪")
        status_label = tk.Label(self.root, textvariable=self.status_var, fg="blue")
        status_label.pack(pady=10)

        self.progress = ttk.Progressbar(self.root, orient="horizontal", length=400, mode="determinate")
        self.progress.pack(pady=10)

        # 执行按钮
        self.btn_run = tk.Button(self.root, text="开始一键转换", bg="green", fg="white", font=("Arial", 12, "bold"), command=self.start_processing)
        self.btn_run.pack(pady=15)

    def browse_file(self):
        filepath = filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")])
        if filepath:
            self.csv_path = filepath
            self.path_label.config(text=os.path.basename(filepath), fg="black")

    def set_cell_font(self, cell, text):
        """核心逻辑：设置单元格文本，并完美分离中英文字体"""
        # 清空单元格默认的段落
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(str(text))
        
        # 统一设置字号
        run.font.size = Pt(11)
        
        # 1. 设置西文字体 (English)
        run.font.name = 'Times New Roman'
        
        # 2. 设置东亚字体 (Chinese) - 需要操作底层 XML
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def process_data(self):
        try:
            self.btn_run.config(state=tk.DISABLED)
            self.status_var.set("正在读取 CSV 数据，请稍候...")
            
            # 读取 CSV
            df = pd.read_csv(self.csv_path)
            total_rows = len(df)
            columns = df.columns.tolist()
            col_count = len(columns)

            # 为了防止 Word 崩溃，每 5000 行切割为一个文件
            chunk_size = 5000
            total_chunks = math.ceil(total_rows / chunk_size)
            
            output_dir = os.path.dirname(self.csv_path)
            base_name = os.path.splitext(os.path.basename(self.csv_path))[0]

            self.progress["maximum"] = total_chunks
            self.progress["value"] = 0

            for i in range(total_chunks):
                self.status_var.set(f"正在生成第 {i+1}/{total_chunks} 个文档...")
                
                start_idx = i * chunk_size
                end_idx = min((i + 1) * chunk_size, total_rows)
                chunk_df = df.iloc[start_idx:end_idx]

                # 创建 Word 文档
                doc = Document()
                
                # 动态生成表格：行数=数据行+1(表头)，列数=CSV列数
                table = doc.add_table(rows=len(chunk_df) + 1, cols=col_count)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # 写入表头并设置字体
                for col_idx, col_name in enumerate(columns):
                    self.set_cell_font(table.cell(0, col_idx), col_name)

                # 写入数据行
                for row_idx, (_, row) in enumerate(chunk_df.iterrows(), start=1):
                    for col_idx in range(col_count):
                        cell_value = row.iloc[col_idx]
                        # 处理空值 (NaN)
                        if pd.isna(cell_value):
                            cell_value = ""
                        self.set_cell_font(table.cell(row_idx, col_idx), cell_value)

                # 保存当前分卷文档
                output_filename = os.path.join(output_dir, f"{base_name}_排版输出_Part{i+1}.docx")
                doc.save(output_filename)
                
                # 更新进度条
                self.progress["value"] = i + 1
                self.root.update_idletasks()

            self.status_var.set("转换彻底完成！请查看 CSV 所在文件夹。")
            messagebox.showinfo("成功", "所有数据已成功转换为带有完美字体的 Word 表格！")

        except Exception as e:
            messagebox.showerror("发生错误", f"处理过程中出现问题：\n{str(e)}")
            self.status_var.set("处理失败")
        finally:
            self.btn_run.config(state=tk.NORMAL)

    def start_processing(self):
        if not self.csv_path:
            messagebox.showwarning("提示", "请先选择一个 CSV 文件！")
            return
        
        # 开启独立线程处理，防止 GUI 界面假死
        thread = threading.Thread(target=self.process_data)
        thread.daemon = True
        thread.start()

if __name__ == "__main__":
    root = tk.Tk()
    app = CSVtoWordApp(root)
    root.mainloop()